import io
import docx  # python-docx for DOCX parsing
from typing import Tuple, Dict
import logging
import re

logger = logging.getLogger(__name__)

# Try to import PyMuPDF, with fallback options
PDF_PARSER = None
try:
    import fitz  # PyMuPDF
    PDF_PARSER = 'pymupdf'
    logger.info("Using PyMuPDF for PDF parsing")
except ImportError:
    try:
        import PyPDF2
        PDF_PARSER = 'pypdf2'
        logger.info("Using PyPDF2 for PDF parsing")
    except ImportError:
        try:
            import pdfplumber
            PDF_PARSER = 'pdfplumber'
            logger.info("Using pdfplumber for PDF parsing")
        except ImportError:
            PDF_PARSER = None
            logger.warning("No PDF parser available. PDF processing will be limited.")

class DocumentParser:
    """Document parser for various file formats with fallback PDF parsing"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.pdf_parser = PDF_PARSER
    
    def parse_pdf(self, content: bytes) -> Tuple[str, Dict]:
        """Parse PDF content and extract text with metadata"""
        try:
            if self.pdf_parser == 'pymupdf':
                return self._parse_pdf_pymupdf(content)
            elif self.pdf_parser == 'pypdf2':
                return self._parse_pdf_pypdf2(content)
            elif self.pdf_parser == 'pdfplumber':
                return self._parse_pdf_pdfplumber(content)
            else:
                # Fallback to basic text extraction attempt
                return self._parse_pdf_fallback(content)
                
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            # Try fallback method
            return self._parse_pdf_fallback(content)
    
    def _parse_pdf_pymupdf(self, content: bytes) -> Tuple[str, Dict]:
        """Parse PDF using PyMuPDF (fitz)"""
        pdf_document = fitz.open(stream=content, filetype="pdf")
        
        text = ""
        metadata = {
            'doc_type': 'pdf',
            'pages': len(pdf_document),
            'title': pdf_document.metadata.get('title', ''),
            'author': pdf_document.metadata.get('author', ''),
            'creator': pdf_document.metadata.get('creator', '')
        }
        
        # Extract text from all pages
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
            text += "\n\n"  # Add page break
        
        pdf_document.close()
        
        # Add word count to metadata
        metadata['word_count'] = len(text.split())
        metadata['char_count'] = len(text)
        
        logger.info(f"PDF parsed with PyMuPDF: {metadata['pages']} pages, {metadata['word_count']} words")
        return text.strip(), metadata
    
    def _parse_pdf_pypdf2(self, content: bytes) -> Tuple[str, Dict]:
        """Parse PDF using PyPDF2"""
        import PyPDF2
        
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        text = ""
        
        metadata = {
            'doc_type': 'pdf',
            'pages': len(pdf_reader.pages),
            'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
            'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
            'creator': pdf_reader.metadata.get('/Creator', '') if pdf_reader.metadata else ''
        }
        
        # Extract text from all pages
        for page_num, page in enumerate(pdf_reader.pages):
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page.extract_text()
            text += "\n\n"
        
        metadata['word_count'] = len(text.split())
        metadata['char_count'] = len(text)
        
        logger.info(f"PDF parsed with PyPDF2: {metadata['pages']} pages, {metadata['word_count']} words")
        return text.strip(), metadata
    
    def _parse_pdf_pdfplumber(self, content: bytes) -> Tuple[str, Dict]:
        """Parse PDF using pdfplumber"""
        import pdfplumber
        
        text = ""
        pages_count = 0
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages_count = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages):
                text += f"\n--- Page {page_num + 1} ---\n"
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                text += "\n\n"
        
        metadata = {
            'doc_type': 'pdf',
            'pages': pages_count,
            'title': '',
            'author': '',
            'creator': 'pdfplumber',
            'word_count': len(text.split()),
            'char_count': len(text)
        }
        
        logger.info(f"PDF parsed with pdfplumber: {metadata['pages']} pages, {metadata['word_count']} words")
        return text.strip(), metadata
    
    def _parse_pdf_fallback(self, content: bytes) -> Tuple[str, Dict]:
        """Fallback PDF parsing - basic text extraction attempt"""
        logger.warning("Using fallback PDF parsing - text extraction may be limited")
        
        # Try to extract readable text using basic string operations
        try:
            # Convert bytes to string and try to find readable text
            content_str = content.decode('latin-1', errors='ignore')
            
            # Simple heuristic to extract text from PDF
            text_parts = []
            lines = content_str.split('\n')
            
            for line in lines:
                # Look for lines that might contain readable text
                if len(line) > 10 and any(c.isalpha() for c in line):
                    # Clean up the line
                    cleaned = re.sub(r'[^\x20-\x7E]', ' ', line)
                    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                    if len(cleaned) > 5:
                        text_parts.append(cleaned)
            
            text = '\n'.join(text_parts)
            
            metadata = {
                'doc_type': 'pdf',
                'pages': 1,  # Can't determine actual pages
                'title': '',
                'author': '',
                'creator': 'fallback_parser',
                'word_count': len(text.split()),
                'char_count': len(text),
                'parsing_method': 'fallback'
            }
            
            logger.info(f"PDF parsed with fallback method: {metadata['word_count']} words extracted")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Fallback PDF parsing also failed: {str(e)}")
            # Return minimal result
            return "PDF content could not be extracted", {
                'doc_type': 'pdf',
                'pages': 1,
                'title': '',
                'author': '',
                'creator': 'failed_parser',
                'word_count': 0,
                'char_count': 0,
                'parsing_method': 'failed'
            }
    
    def parse_pdf_enhanced(self, content: bytes) -> Tuple[str, Dict]:
        """Enhanced PDF parsing - SYNCHRONOUS METHOD"""
        try:
            if self.pdf_parser == 'pymupdf':
                return self._parse_pdf_enhanced_pymupdf(content)
            else:
                # Use regular parsing for other parsers
                return self.parse_pdf(content)
                
        except Exception as e:
            logger.error(f"Enhanced PDF parsing failed: {str(e)}")
            # Fallback to basic parsing
            return self.parse_pdf(content)
    
    def _parse_pdf_enhanced_pymupdf(self, content: bytes) -> Tuple[str, Dict]:
        """Enhanced PDF parsing using PyMuPDF"""
        pdf_document = fitz.open(stream=content, filetype="pdf")
        
        text = ""
        metadata = {
            'doc_type': 'pdf',
            'pages': len(pdf_document),
            'title': pdf_document.metadata.get('title', ''),
            'author': pdf_document.metadata.get('author', ''),
            'creator': pdf_document.metadata.get('creator', ''),
            'page_texts': []
        }
        
        # Enhanced text extraction with page-level metadata
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            page_text = page.get_text("text")  # More reliable text extraction
            
            # Store page-level information
            metadata['page_texts'].append({
                'page_num': page_num + 1,
                'text_length': len(page_text),
                'word_count': len(page_text.split())
            })
            
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page_text
            text += "\n\n"
        
        pdf_document.close()
        
        # Add enhanced metadata
        metadata['word_count'] = len(text.split())
        metadata['char_count'] = len(text)
        metadata['avg_words_per_page'] = metadata['word_count'] / max(1, metadata['pages'])
        
        logger.info(f"Enhanced PDF parsed: {metadata['pages']} pages, {metadata['word_count']} words")
        return text.strip(), metadata
    
    def parse_docx(self, content: bytes) -> Tuple[str, Dict]:
        """Parse DOCX content and extract text with metadata"""
        try:
            # Open DOCX from bytes
            doc = docx.Document(io.BytesIO(content))
            
            text = ""
            metadata = {
                'doc_type': 'docx',
                'paragraphs': 0,
                'tables': 0
            }
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    metadata['paragraphs'] += 1
            
            # Extract text from tables
            for table in doc.tables:
                metadata['tables'] += 1
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            # Add word count to metadata
            metadata['word_count'] = len(text.split())
            metadata['char_count'] = len(text)
            
            logger.info(f"DOCX parsed successfully: {metadata['paragraphs']} paragraphs, {metadata['word_count']} words")
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            raise
    
    def parse_docx_enhanced(self, content: bytes) -> Tuple[str, Dict]:
        """Enhanced DOCX parsing - SYNCHRONOUS METHOD"""
        try:
            doc = docx.Document(io.BytesIO(content))
            
            text = ""
            metadata = {
                'doc_type': 'docx',
                'paragraphs': 0,
                'tables': 0,
                'headers': 0,
                'sections': len(doc.sections),
                'styles_used': set()
            }
            
            # Enhanced paragraph extraction with style information
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Track styles
                    if paragraph.style:
                        metadata['styles_used'].add(paragraph.style.name)
                    
                    # Check if it's likely a header
                    if paragraph.style and ('head' in paragraph.style.name.lower() or 
                                          'title' in paragraph.style.name.lower()):
                        metadata['headers'] += 1
                        text += f"\n## {paragraph.text}\n"
                    else:
                        text += paragraph.text + "\n"
                    
                    metadata['paragraphs'] += 1
            
            # Enhanced table extraction
            for table in doc.tables:
                metadata['tables'] += 1
                text += "\n--- Table ---\n"
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "--- End Table ---\n\n"
            
            # Convert styles_used set to list for JSON serialization
            metadata['styles_used'] = list(metadata['styles_used'])
            
            # Add enhanced metadata
            metadata['word_count'] = len(text.split())
            metadata['char_count'] = len(text)
            metadata['avg_words_per_paragraph'] = metadata['word_count'] / max(1, metadata['paragraphs'])
            
            logger.info(f"Enhanced DOCX parsed: {metadata['paragraphs']} paragraphs, {metadata['headers']} headers")
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Enhanced DOCX parsing failed: {str(e)}")
            # Fallback to basic parsing
            return self.parse_docx(content)
    
    def parse_text(self, content: bytes) -> Tuple[str, Dict]:
        """Parse plain text content"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            metadata = {
                'doc_type': 'txt',
                'word_count': len(text.split()),
                'char_count': len(text),
                'lines': len(text.split('\n'))
            }
            
            logger.info(f"Text parsed successfully: {metadata['lines']} lines, {metadata['word_count']} words")
            return text, metadata
            
        except Exception as e:
            logger.error(f"Text parsing failed: {str(e)}")
            raise
    
    def is_supported_format(self, file_extension: str) -> bool:
        """Check if file format is supported"""
        return file_extension.lower() in self.supported_formats